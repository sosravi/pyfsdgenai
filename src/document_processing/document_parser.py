"""
Document Parser

This module implements the DocumentParser class for parsing various document formats
and extracting text content and metadata.
"""

import os
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parses various document formats and extracts text content."""
    
    def __init__(self, max_text_length: int = 1000000):
        """
        Initialize DocumentParser.
        
        Args:
            max_text_length: Maximum text length to extract (default: 1MB)
        """
        self.supported_formats = ["pdf", "doc", "docx", "txt"]
        self.max_text_length = max_text_length
        
        logger.info("DocumentParser initialized")
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing parsed content and metadata
        """
        try:
            file_path = Path(file_path)
            
            # Validate file exists
            if not file_path.exists():
                return {
                    "success": False,
                    "error": "File does not exist"
                }
            
            # Get file extension
            file_extension = file_path.suffix.lower().lstrip('.')
            
            # Validate format
            if file_extension not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"Unsupported file format: {file_extension}. Supported formats: {', '.join(self.supported_formats)}"
                }
            
            # Parse based on format
            if file_extension == "pdf":
                return self._parse_pdf(file_path)
            elif file_extension == "docx":
                return self._parse_docx(file_path)
            elif file_extension == "doc":
                return self._parse_doc(file_path)
            elif file_extension == "txt":
                return self._parse_txt(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Parser not implemented for format: {file_extension}"
                }
                
        except Exception as e:
            logger.error(f"Error parsing document: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to parse document: {str(e)}"
            }
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF document."""
        try:
            # Try to import PyPDF2
            try:
                import PyPDF2
            except ImportError:
                # Fallback to simple text extraction
                return self._parse_txt(file_path)
            
            text_content = ""
            metadata = {
                "format": "pdf",
                "pages": 0,
                "word_count": 0
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    text_content += page_text + "\n"
                    
                    # Check text length limit
                    if len(text_content) > self.max_text_length:
                        text_content = text_content[:self.max_text_length]
                        logger.warning(f"Text content truncated to {self.max_text_length} characters")
                        break
            
            metadata["word_count"] = len(text_content.split())
            
            return {
                "success": True,
                "text_content": text_content.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to parse PDF: {str(e)}"
            }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX document."""
        try:
            # Try to import python-docx
            try:
                from docx import Document
            except ImportError:
                # Fallback to simple text extraction
                return self._parse_txt(file_path)
            
            text_content = ""
            metadata = {
                "format": "docx",
                "paragraphs": 0,
                "word_count": 0
            }
            
            doc = Document(file_path)
            metadata["paragraphs"] = len(doc.paragraphs)
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
                
                # Check text length limit
                if len(text_content) > self.max_text_length:
                    text_content = text_content[:self.max_text_length]
                    logger.warning(f"Text content truncated to {self.max_text_length} characters")
                    break
            
            metadata["word_count"] = len(text_content.split())
            
            return {
                "success": True,
                "text_content": text_content.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to parse DOCX: {str(e)}"
            }
    
    def _parse_doc(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOC document."""
        try:
            # DOC parsing is complex, for now return error
            # In production, would use python-docx2txt or similar
            return {
                "success": False,
                "error": "DOC format parsing not implemented. Please convert to DOCX or PDF."
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOC: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to parse DOC: {str(e)}"
            }
    
    def _parse_txt(self, file_path: Path) -> Dict[str, Any]:
        """Parse TXT document."""
        try:
            text_content = ""
            metadata = {
                "format": "txt",
                "lines": 0,
                "word_count": 0
            }
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                lines = file.readlines()
                metadata["lines"] = len(lines)
                
                for line in lines:
                    text_content += line
                    
                    # Check text length limit
                    if len(text_content) > self.max_text_length:
                        text_content = text_content[:self.max_text_length]
                        logger.warning(f"Text content truncated to {self.max_text_length} characters")
                        break
            
            metadata["word_count"] = len(text_content.split())
            
            return {
                "success": True,
                "text_content": text_content.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing TXT: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to parse TXT: {str(e)}"
            }
    
    def extract_text_sections(self, text: str) -> Dict[str, str]:
        """
        Extract structured sections from text content.
        
        Args:
            text: Text content to analyze
            
        Returns:
            Dict containing extracted sections
        """
        try:
            sections = {
                "parties": "",
                "terms": "",
                "pricing": "",
                "signatures": ""
            }
            
            # Extract parties
            parties_patterns = [
                r'parties?[:\s]*([^\n]+)',
                r'between[:\s]*([^\n]+)',
                r'company[:\s]*([^\n]+)',
            ]
            
            for pattern in parties_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    sections["parties"] = matches[0].strip()
                    break
            
            # Extract terms
            terms_patterns = [
                r'terms?[:\s]*([^\n]+)',
                r'conditions?[:\s]*([^\n]+)',
                r'payment[:\s]*([^\n]+)',
            ]
            
            for pattern in terms_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    sections["terms"] = matches[0].strip()
                    break
            
            # Extract pricing
            pricing_patterns = [
                r'price[:\s]*([^\n]+)',
                r'cost[:\s]*([^\n]+)',
                r'amount[:\s]*([^\n]+)',
                r'\$[\d,]+\.?\d*',
            ]
            
            for pattern in pricing_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    sections["pricing"] = matches[0].strip()
                    break
            
            # Extract signatures
            signature_patterns = [
                r'signature[:\s]*([^\n]+)',
                r'signed[:\s]*([^\n]+)',
                r'authorized[:\s]*([^\n]+)',
            ]
            
            for pattern in signature_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    sections["signatures"] = matches[0].strip()
                    break
            
            return sections
            
        except Exception as e:
            logger.error(f"Error extracting text sections: {str(e)}")
            return {
                "parties": "",
                "terms": "",
                "pricing": "",
                "signatures": ""
            }
    
    def get_text_statistics(self, text: str) -> Dict[str, Any]:
        """
        Get statistics about the text content.
        
        Args:
            text: Text content to analyze
            
        Returns:
            Dict containing text statistics
        """
        try:
            words = text.split()
            sentences = re.split(r'[.!?]+', text)
            paragraphs = text.split('\n\n')
            
            return {
                "character_count": len(text),
                "word_count": len(words),
                "sentence_count": len([s for s in sentences if s.strip()]),
                "paragraph_count": len([p for p in paragraphs if p.strip()]),
                "average_word_length": sum(len(word) for word in words) / len(words) if words else 0,
                "average_sentence_length": len(words) / len([s for s in sentences if s.strip()]) if sentences else 0
            }
            
        except Exception as e:
            logger.error(f"Error calculating text statistics: {str(e)}")
            return {
                "character_count": 0,
                "word_count": 0,
                "sentence_count": 0,
                "paragraph_count": 0,
                "average_word_length": 0,
                "average_sentence_length": 0
            }
